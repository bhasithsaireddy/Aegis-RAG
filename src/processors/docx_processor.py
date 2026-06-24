"""Advanced DOCX document processor with structure-aware extraction"""
from pathlib import Path
from typing import List, Optional, Dict, Any
import tempfile
import logging

try:
    from docx import Document
    from docx.table import Table
except ImportError:
    Document = None
    Table = None

from .base import BaseProcessor, Chunk
from ..chunking import Chunker
from ..config import config

logger = logging.getLogger(__name__)


class DOCXProcessor(BaseProcessor):
    """
    Advanced DOCX processor with:
    - Section/heading-aware chunking
    - Table extraction to markdown
    - Embedded image extraction with OCR
    - Style and formatting metadata
    - Hyperlink preservation
    """
    
    @property
    def supported_extensions(self) -> List[str]:
        return [".docx", ".doc"]
    
    @property
    def doc_type(self) -> str:
        return "docx"
    
    def __init__(
        self,
        chunker: Optional[Chunker] = None,
        use_ocr_for_images: bool = True,
        preserve_sections: bool = True
    ):
        """
        Initialize DOCX processor.
        
        Args:
            chunker: Text chunker instance
            use_ocr_for_images: Whether to OCR embedded images
            preserve_sections: Whether to chunk by section headings
        """
        if Document is None:
            raise ImportError("python-docx not installed. Run: pip install python-docx")
        
        self.chunker = chunker or Chunker()
        self.use_ocr_for_images = use_ocr_for_images
        self.preserve_sections = preserve_sections
        self._ocr = None
    
    def _get_ocr(self):
        """Lazy load OCR engine"""
        if self._ocr is None and self.use_ocr_for_images:
            try:
                from ..ocr import DeepSeekOCR
                self._ocr = DeepSeekOCR()
                if not self._ocr.is_available():
                    logger.warning("DeepSeek model not found, image OCR disabled")
                    self._ocr = None
            except Exception as e:
                logger.warning(f"OCR initialization failed: {e}")
                self._ocr = None
        return self._ocr
    
    def process(self, file_path: Path, document_id: Optional[str] = None) -> List[Chunk]:
        """Process a DOCX file and extract chunks."""
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"DOCX file not found: {file_path}")
        
        document_id = document_id or self._generate_document_id(file_path)
        chunks: List[Chunk] = []
        
        doc = Document(file_path)
        
        # Extract document metadata
        doc_metadata = self._extract_metadata(doc, file_path)
        
        # Process paragraphs with section awareness
        if self.preserve_sections:
            chunks.extend(self._process_with_sections(doc, document_id, file_path, doc_metadata))
        else:
            chunks.extend(self._process_flat(doc, document_id, file_path, doc_metadata))
        
        # Extract tables
        for table_idx, table in enumerate(doc.tables):
            table_md = self._table_to_markdown(table)
            if table_md:
                chunks.append(Chunk(
                    content=f"[Table {table_idx + 1}]\n{table_md}",
                    document_id=document_id,
                    doc_type=self.doc_type,
                    source_file=str(file_path),
                    metadata={
                        **doc_metadata,
                        "content_type": "table",
                        "table_index": table_idx,
                        "rows": len(table.rows),
                        "columns": len(table.columns)
                    }
                ))
        
        # Extract embedded images
        if self.use_ocr_for_images:
            image_chunks = self._extract_embedded_images(file_path, doc, document_id, doc_metadata)
            chunks.extend(image_chunks)
        
        return chunks
    
    def _extract_metadata(self, doc: Document, file_path: Path) -> Dict[str, Any]:
        """Extract document metadata."""
        core = doc.core_properties
        return {
            "title": core.title or file_path.stem,
            "author": core.author or "",
            "subject": core.subject or "",
            "keywords": core.keywords or "",
            "created": str(core.created) if core.created else "",
            "modified": str(core.modified) if core.modified else "",
        }
    
    def _process_with_sections(
        self, 
        doc: Document, 
        document_id: str, 
        file_path: Path,
        doc_metadata: Dict[str, Any]
    ) -> List[Chunk]:
        """Process document with section/heading awareness."""
        chunks = []
        current_section = ""
        section_text = []
        section_level = 0
        
        for para in doc.paragraphs:
            style_name = para.style.name if para.style else ""
            
            # Check if heading
            if style_name.startswith('Heading'):
                # Save previous section
                if section_text:
                    chunks.extend(self._chunk_section(
                        "\n".join(section_text),
                        current_section,
                        section_level,
                        document_id,
                        file_path,
                        doc_metadata
                    ))
                    section_text = []
                
                # Parse heading level
                try:
                    section_level = int(style_name.replace('Heading ', '').strip())
                except ValueError:
                    section_level = 1
                
                current_section = para.text.strip()
            else:
                if para.text.strip():
                    section_text.append(para.text)
        
        # Don't forget the last section
        if section_text:
            chunks.extend(self._chunk_section(
                "\n".join(section_text),
                current_section,
                section_level,
                document_id,
                file_path,
                doc_metadata
            ))
        
        return chunks
    
    def _chunk_section(
        self,
        text: str,
        section_title: str,
        section_level: int,
        document_id: str,
        file_path: Path,
        doc_metadata: Dict[str, Any]
    ) -> List[Chunk]:
        """Chunk a section and add metadata."""
        chunks = []
        text_chunks = self.chunker.chunk(text)
        
        for idx, chunk_text in enumerate(text_chunks):
            content = chunk_text
            if section_title:
                content = f"[Section: {section_title}]\n\n{chunk_text}"
            
            chunks.append(Chunk(
                content=content,
                document_id=document_id,
                doc_type=self.doc_type,
                source_file=str(file_path),
                metadata={
                    **doc_metadata,
                    "section": section_title,
                    "section_level": section_level,
                    "chunk_index": idx,
                    "content_type": "text"
                }
            ))
        
        return chunks
    
    def _process_flat(
        self, 
        doc: Document, 
        document_id: str, 
        file_path: Path,
        doc_metadata: Dict[str, Any]
    ) -> List[Chunk]:
        """Process document without section awareness."""
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        
        chunks = []
        text_chunks = self.chunker.chunk("\n".join(full_text))
        
        for idx, chunk_text in enumerate(text_chunks):
            chunks.append(Chunk(
                content=chunk_text,
                document_id=document_id,
                doc_type=self.doc_type,
                source_file=str(file_path),
                metadata={
                    **doc_metadata,
                    "chunk_index": idx,
                    "content_type": "text"
                }
            ))
        
        return chunks
    
    def _extract_embedded_images(
        self, 
        file_path: Path, 
        doc: Document, 
        document_id: str,
        doc_metadata: Dict[str, Any]
    ) -> List[Chunk]:
        """Extract and OCR embedded images."""
        chunks = []
        ocr = self._get_ocr()
        
        if ocr is None:
            return chunks
        
        try:
            img_idx = 0
            for rel_id, rel in doc.part.rels.items():
                if "image" in rel.target_ref:
                    try:
                        image_data = rel.target_part.blob
                        
                        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
                            tmp_path = Path(tmp.name)
                            tmp.write(image_data)
                        
                        try:
                            # Extract text from image
                            content = ocr.extract_full_content(tmp_path)
                            
                            if content.get("text", "").strip():
                                chunks.append(Chunk(
                                    content=f"[Embedded Image Text]\n{content['text']}",
                                    document_id=document_id,
                                    doc_type=self.doc_type,
                                    source_file=str(file_path),
                                    metadata={
                                        **doc_metadata,
                                        "content_type": "embedded_image_text",
                                        "image_index": img_idx,
                                        "extraction_method": "ocr"
                                    }
                                ))
                            
                            # Visual description
                            description = ocr.get_visual_description(tmp_path)
                            if description:
                                chunks.append(Chunk(
                                    content=f"[Image Description]\n{description}",
                                    document_id=document_id,
                                    doc_type=self.doc_type,
                                    source_file=str(file_path),
                                    metadata={
                                        **doc_metadata,
                                        "content_type": "image_description",
                                        "image_index": img_idx,
                                        "extraction_method": "vision"
                                    }
                                ))
                            
                            # Tables from image
                            for table_idx, table_md in enumerate(content.get("tables", [])):
                                chunks.append(Chunk(
                                    content=f"[Table from Image]\n{table_md}",
                                    document_id=document_id,
                                    doc_type=self.doc_type,
                                    source_file=str(file_path),
                                    metadata={
                                        **doc_metadata,
                                        "content_type": "image_table",
                                        "table_index": table_idx,
                                        "image_index": img_idx
                                    }
                                ))
                        finally:
                            if tmp_path.exists():
                                tmp_path.unlink()
                        
                        img_idx += 1
                    except Exception as e:
                        logger.debug(f"Failed to process image: {e}")
        except Exception as e:
            logger.debug(f"Failed to extract images: {e}")
        
        return chunks
    
    def _table_to_markdown(self, table: Table) -> str:
        """Convert DOCX table to markdown format."""
        rows = []
        
        for row in table.rows:
            cells = [cell.text.strip().replace("|", "\\|") for cell in row.cells]
            rows.append("| " + " | ".join(cells) + " |")
        
        if len(rows) >= 1:
            num_cols = len(table.rows[0].cells)
            header_sep = "| " + " | ".join(["---"] * num_cols) + " |"
            rows.insert(1, header_sep)
        
        return "\n".join(rows) if rows else ""
